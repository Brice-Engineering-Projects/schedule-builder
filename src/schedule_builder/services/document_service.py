"""Document upload and text extraction workflow."""

from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

import fitz  # type: ignore[import-untyped]
import pdfplumber  # type: ignore[import-untyped]
from docx import Document as DocxDocument  # type: ignore[import-untyped]

from schedule_builder.config.settings import settings
from schedule_builder.core.exceptions import BadRequestError, NotFoundError, ValidationFailedError
from schedule_builder.models.document import Document
from schedule_builder.repositories.document_repository import DocumentRepository
from schedule_builder.repositories.project_repository import ProjectRepository
from schedule_builder.schemas.document import ExtractedDocumentText


class DocumentService:
    """Orchestrates document validation, storage, and extraction."""

    allowed_extensions = {"pdf", "docx", "txt"}

    def __init__(
        self,
        document_repository: DocumentRepository,
        project_repository: ProjectRepository,
        upload_directory: str,
    ) -> None:
        self._documents = document_repository
        self._projects = project_repository
        self._upload_directory = Path(upload_directory)
        self._max_upload_bytes = settings.max_upload_size_mb * 1024 * 1024
        self._minimum_text_length = 25

    def ingest_document(
        self,
        *,
        project_id: str,
        filename: str,
        content_type: str,
        file_bytes: bytes,
    ) -> tuple[Document, ExtractedDocumentText]:
        project = self._projects.get_by_id(project_id)
        if project is None:
            raise NotFoundError(message="Project not found", details={"project_id": project_id})

        cleaned_filename = Path(filename).name
        file_extension = self._get_file_extension(cleaned_filename)
        self._validate_extension(file_extension)
        self._validate_file_size(file_bytes)

        document_id = str(uuid4())
        storage_path = self._build_storage_path(project_id, document_id, cleaned_filename)
        document = self._documents.create(
            document_id=document_id,
            project_id=project_id,
            filename=cleaned_filename,
            content_type=content_type,
            file_size_bytes=len(file_bytes),
            storage_path=str(storage_path),
            processing_status="processing",
        )

        try:
            self._write_file(storage_path, file_bytes)
            extracted_text = self.extract_document_text(
                filename=cleaned_filename,
                content_type=content_type,
                file_bytes=file_bytes,
            )
            if len(extracted_text.text.strip()) < self._minimum_text_length:
                raise ValidationFailedError(
                    message="Extracted text is too short to process",
                    details={
                        "document_id": document.id,
                        "character_count": extracted_text.character_count,
                    },
                )

            extracted_text = extracted_text.model_copy(update={"document_id": document.id})

            document = self._documents.update(
                document,
                storage_path=str(storage_path),
                processing_status="extracted",
                extracted_text=extracted_text.text,
                page_count=extracted_text.page_count,
                extracted_at=extracted_text.extracted_at,
                extraction_error=None,
            )
            return document, extracted_text
        except Exception as exc:
            self._documents.update(
                document,
                processing_status="failed",
                extraction_error=str(exc),
            )
            if storage_path.exists():
                storage_path.unlink()
            if isinstance(exc, ValidationFailedError):
                raise
            raise BadRequestError(
                message="Failed to process uploaded document",
                details={"document_id": document.id, "error": str(exc)},
            ) from exc

    def get_document(self, document_id: str) -> Document:
        document = self._documents.get_by_id(document_id)
        if document is None:
            raise NotFoundError(message="Document not found", details={"document_id": document_id})
        return document

    def get_document_text(self, document_id: str) -> ExtractedDocumentText:
        document = self.get_document(document_id)
        if not document.extracted_text:
            raise ValidationFailedError(
                message="Document text has not been extracted yet",
                details={
                    "document_id": document_id,
                    "processing_status": document.processing_status,
                },
            )

        return ExtractedDocumentText(
            document_id=document.id,
            source_filename=document.filename,
            content_type=document.content_type,
            text=document.extracted_text,
            character_count=len(document.extracted_text),
            page_count=document.page_count,
            extracted_at=document.extracted_at or document.updated_at,
        )

    def extract_document_text(
        self,
        *,
        filename: str,
        content_type: str,
        file_bytes: bytes,
    ) -> ExtractedDocumentText:
        file_extension = self._get_file_extension(filename)

        if file_extension == "pdf":
            text, page_count = self._extract_pdf_text(file_bytes)
        elif file_extension == "docx":
            text, page_count = self._extract_docx_text(file_bytes)
        elif file_extension == "txt":
            text, page_count = self._extract_txt_text(file_bytes)
        else:
            raise BadRequestError(
                message="Unsupported document type",
                details={"filename": filename, "extension": file_extension},
            )

        normalized_text = self._normalize_text(text)
        return ExtractedDocumentText(
            document_id="",
            source_filename=filename,
            content_type=content_type,
            text=normalized_text,
            character_count=len(normalized_text),
            page_count=page_count,
            extracted_at=datetime.now(tz=timezone.utc),
        )

    def _extract_pdf_text(self, file_bytes: bytes) -> tuple[str, int | None]:
        try:
            with fitz.open(stream=file_bytes, filetype="pdf") as pdf_document:
                text = "\n".join(page.get_text("text") for page in pdf_document)
                page_count = pdf_document.page_count
        except Exception:
            text = ""
            page_count = None

        if not text.strip():
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf_document:
                text = "\n".join(page.extract_text() or "" for page in pdf_document.pages)
                page_count = len(pdf_document.pages)

        return text, page_count

    def _extract_docx_text(self, file_bytes: bytes) -> tuple[str, int | None]:
        document = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_rows: list[str] = []

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_rows.append(" | ".join(cells))

        text = "\n".join(paragraphs + table_rows)
        return text, None

    def _extract_txt_text(self, file_bytes: bytes) -> tuple[str, int | None]:
        for encoding in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                return file_bytes.decode(encoding), 1
            except UnicodeDecodeError:
                continue

        raise BadRequestError(message="Unable to decode text document", details={})

    def _build_storage_path(self, project_id: str, document_id: str, filename: str) -> Path:
        return self._upload_directory / project_id / document_id / filename

    @staticmethod
    def _get_file_extension(filename: str) -> str:
        return Path(filename).suffix.lower().lstrip(".")

    def _validate_extension(self, file_extension: str) -> None:
        if file_extension not in self.allowed_extensions:
            raise BadRequestError(
                message="Unsupported file type",
                details={
                    "allowed_extensions": sorted(self.allowed_extensions),
                    "extension": file_extension,
                },
            )

    def _validate_file_size(self, file_bytes: bytes) -> None:
        if len(file_bytes) > self._max_upload_bytes:
            raise BadRequestError(
                message="Uploaded file exceeds the maximum allowed size",
                details={
                    "max_upload_bytes": self._max_upload_bytes,
                    "file_size_bytes": len(file_bytes),
                },
            )

    def _write_file(self, storage_path: Path, file_bytes: bytes) -> None:
        storage_path.parent.mkdir(parents=True, exist_ok=True)
        storage_path.write_bytes(file_bytes)

    @staticmethod
    def _normalize_text(text: str) -> str:
        cleaned = text.replace("\r\n", "\n").replace("\r", "\n")
        cleaned = re.sub(r"[\t ]+\n", "\n", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned.strip()
