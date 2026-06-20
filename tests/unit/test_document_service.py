from __future__ import annotations

import io

import fitz  # type: ignore[import-untyped]
from docx import Document as DocxDocument  # type: ignore[import-untyped]

from schedule_builder.models.project import Project
from schedule_builder.models.user import User
from schedule_builder.repositories.document_repository import DocumentRepository
from schedule_builder.repositories.project_repository import ProjectRepository
from schedule_builder.services.document_service import DocumentService


def _build_pdf_bytes(text: str) -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    return document.tobytes()


def _build_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_paragraph("Project Overview")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Discipline"
    table.cell(0, 1).text = "Deliverable"
    table.cell(1, 0).text = "Civil"
    table.cell(1, 1).text = "Site Plan"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_txt_bytes(text: str) -> bytes:
    return text.encode("utf-8")


def _build_service(db_session, upload_dir: str = "uploads") -> DocumentService:
    return DocumentService(
        DocumentRepository(db_session),
        ProjectRepository(db_session),
        upload_dir,
    )


def _create_project(db_session) -> str:
    owner = User(email="owner@example.com", full_name="Owner", password_hash="hash")
    db_session.add(owner)
    db_session.flush()
    project = Project(owner_user_id=owner.id, name="Test Project", description=None)
    db_session.add(project)
    db_session.commit()
    return project.id


def test_extract_pdf_text(db_session) -> None:
    service = _build_service(db_session)

    result = service.extract_document_text(
        filename="sample.pdf",
        content_type="application/pdf",
        file_bytes=_build_pdf_bytes("Hello PDF text"),
    )

    assert "Hello PDF text" in result.text
    assert result.page_count == 1


def test_extract_docx_text(db_session) -> None:
    service = _build_service(db_session)

    result = service.extract_document_text(
        filename="sample.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_bytes=_build_docx_bytes(),
    )

    assert "Project Overview" in result.text
    assert "Civil" in result.text
    assert result.page_count is None


def test_extract_txt_text(db_session) -> None:
    service = _build_service(db_session)

    result = service.extract_document_text(
        filename="sample.txt",
        content_type="text/plain",
        file_bytes=_build_txt_bytes("Simple scope text"),
    )

    assert result.text == "Simple scope text"
    assert result.page_count == 1


def test_ingest_document_persists_record_and_text(db_session, tmp_path) -> None:
    project_id = _create_project(db_session)
    service = _build_service(db_session, upload_dir=str(tmp_path / "uploads"))

    document, extracted_text = service.ingest_document(
        project_id=project_id,
        filename="scope.pdf",
        content_type="application/pdf",
        file_bytes=_build_pdf_bytes("Civil engineering scope document"),
    )

    assert document.project_id == project_id
    assert document.processing_status == "extracted"
    assert document.extracted_text is not None
    assert extracted_text.document_id == document.id
    assert (tmp_path / "uploads" / project_id / document.id / "scope.pdf").exists()
